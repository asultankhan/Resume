import json
import re

import streamlit as st
from docx import Document
from google import genai
from pypdf import PdfReader


st.set_page_config(page_title="Resume ATS Analyzer", page_icon="📄", layout="wide")


def extract_resume_text(uploaded_file):
    """Extract text from an uploaded PDF or DOCX resume."""
    extension = uploaded_file.name.rsplit(".", 1)[-1].lower()

    if extension == "pdf":
        reader = PdfReader(uploaded_file)
        return "\n".join(page.extract_text() or "" for page in reader.pages).strip()

    if extension == "docx":
        document = Document(uploaded_file)
        paragraphs = [p.text for p in document.paragraphs]
        table_text = [
            cell.text
            for table in document.tables
            for row in table.rows
            for cell in row.cells
        ]
        return "\n".join(paragraphs + table_text).strip()

    raise ValueError("Please upload a PDF or DOCX file.")


def parse_json_response(response_text):
    """Parse JSON even if the model wraps it in a Markdown code block."""
    cleaned = re.sub(r"^```(?:json)?\s*|\s*```$", "", response_text.strip(), flags=re.I)
    start, end = cleaned.find("{"), cleaned.rfind("}")
    if start == -1 or end == -1:
        raise ValueError("Gemini did not return a valid analysis.")
    return json.loads(cleaned[start : end + 1])


def analyze_resume(resume_text, job_description, api_key):
    client = genai.Client(api_key=api_key)
    prompt = f"""
You are an expert ATS resume evaluator and career coach. Compare the resume with
the job description. Base the score only on evidence found in the supplied text.
Do not invent qualifications. Return ONLY valid JSON with exactly this structure:
{{
  "ats_score": 0,
  "overall_assessment": "brief assessment",
  "matched_keywords": ["keyword"],
  "missing_keywords": ["keyword"],
  "strengths": ["strength"],
  "improvements": ["specific improvement"],
  "suggested_summary": "an improved professional summary tailored to the job"
}}

Scoring guidance: skills/keywords 40%, relevant experience 30%, education and
certifications 15%, measurable achievements 10%, ATS readability 5%.

RESUME:
{resume_text}

JOB DESCRIPTION:
{job_description}
"""
    response = client.models.generate_content(
        model="gemini-2.5-flash",
        contents=prompt,
        config={"response_mime_type": "application/json", "temperature": 0.2},
    )
    return parse_json_response(response.text)


st.title("📄 Resume ATS Analyzer")
st.caption("Upload your resume, add a job description, and receive an AI-assisted ATS review.")

with st.sidebar:
    st.header("How it works")
    st.write("1. Upload a PDF or DOCX resume.")
    st.write("2. Paste the target job description.")
    st.write("3. Select **Analyze Resume**.")
    st.info("AI feedback is advisory. Actual ATS systems and recruiters may evaluate resumes differently.")

uploaded_file = st.file_uploader("Upload your resume", type=["pdf", "docx"])
job_description = st.text_area(
    "Paste the job description",
    height=220,
    placeholder="Paste the complete vacancy or job description here...",
)

if st.button("Analyze Resume", type="primary", use_container_width=True):
    if uploaded_file is None:
        st.warning("Please upload your resume.")
    elif not job_description.strip():
        st.warning("Please paste the job description.")
    elif "GEMINI_API_KEY" not in st.secrets:
        st.error("GEMINI_API_KEY is missing. Add it in Streamlit → App settings → Secrets.")
    else:
        try:
            with st.spinner("Analyzing your resume..."):
                resume_text = extract_resume_text(uploaded_file)
                if len(resume_text) < 50:
                    raise ValueError(
                        "Very little text was found. Please upload a text-based PDF or DOCX file."
                    )
                result = analyze_resume(
                    resume_text,
                    job_description.strip(),
                    st.secrets["GEMINI_API_KEY"],
                )

            score = max(0, min(100, int(result.get("ats_score", 0))))
            st.subheader("Your Results")
            left, right = st.columns([1, 2])
            with left:
                st.metric("Estimated ATS Score", f"{score}/100")
                st.progress(score)
            with right:
                st.write(result.get("overall_assessment", "Analysis completed."))

            col1, col2 = st.columns(2)
            with col1:
                st.subheader("✅ Matched Keywords")
                st.write(", ".join(result.get("matched_keywords", [])) or "None identified")
                st.subheader("💪 Strengths")
                for item in result.get("strengths", []):
                    st.write(f"• {item}")
            with col2:
                st.subheader("🔎 Missing Keywords")
                st.write(", ".join(result.get("missing_keywords", [])) or "None identified")
                st.subheader("🛠️ Recommended Improvements")
                for item in result.get("improvements", []):
                    st.write(f"• {item}")

            st.subheader("Suggested Professional Summary")
            st.write(result.get("suggested_summary", "No summary generated."))

        except Exception as error:
            st.error(f"The analysis could not be completed: {error}")

